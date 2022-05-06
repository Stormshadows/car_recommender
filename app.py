from flask import Flask, render_template, request
from flask import send_file
from docx import Document
from io import BytesIO
import psycopg2,logging

dbcon = psycopg2.connect(
        user='rwozksuc',
        password='3nv4b-4aaJb5bx0--2hIAZeoYVXateTm',
        database='rwozksuc',
        host='john.db.elephantsql.com')

app = Flask(__name__)

@app.route("/")
def form():
    return render_template('Car_search.html')

@app.route("/download")
def download_doc():
    
    car = (3, 'Hyundai Creta', 'Hyundai', 'https://imgd.aeplcdn.com/664x374/n/cw/ec/41564/hyundai-creta-right-front-three-quarter9.jpeg?q=85', 'The Hyundai Creta has now entered its second generation and was first showcased at the 2020 Auto Expo. It’s the king of the D-segment and is offered in 14 variants across three engine and gearbox options.\xa0On the outside, it gets distinctive highlights in the form of three dimensional cascading grille with muscle like vertical and horizontal patterns and sculpted bumper with skid plates. The flared wheel arches with thoughtfully crafted crease lines enhance its overall appeal.\xa0The vehicle gets two boomerang shaped LED DRLs, LED Head Lamps and smartly positioned fog lamps to give the SUV a futuristic look. The new Creta gets class-leading features like panoramic sunroof, LED headlamps and enhanced Bluelink 1.5. The unveiled vehicle featured blacked-out windows and windscreens and more details about the interior will be known closer to its launch date.\xa0The Creta is offered with BS6 compliant petrol and diesel engine options. The standard petrol motor is a 1.5-litre producing 112bhp/144Nm and can be had with a CVT or a six-speed manual. The more exciting petrol motor is a 1.4-litre Turbo GDi unit producing 138bhp/242Nm and is only being offered with a seven-speed DCT. The sole diesel on offer is a 1.5-litre four-cylinder unit producing 113bhp/250Nm and can be had with a six-speed manual or a six-speed automatic. \xa0\xa0The Hyundai Creta is the most successful car in the segment since it was launched in 2015 and took the segment by a storm thanks to its combination of powerful engines and strong feature list. It’s a rival to the likes of the Kia Seltos, Nissan Kicks, Renault Duster and Captur as well as cars like the MG Hector and Tata Harrier.', 10.22, 17.93, 17.00, 21.00, True, True, True, True, False, False, 5)
    document = Document()
    head =document.add_heading(car[1],0)
    head.alignment = 1
    document.add_paragraph(str(car[4]))
    document.add_paragraph('Price range: '+str(car[5])+' lakh to '+str(car[6])+' lakh')
    document.add_paragraph('Mileage: '+str(car[7])+' kmpl '+'-'+str(car[8])+' kmpl')
    document.add_paragraph('Seating capacity: '+str(car[15]))
    f = BytesIO()
    # do staff with document
    document.save(f)
    f.seek(0)

    return send_file(
        f,
        as_attachment=True,
        attachment_filename='car_doc.docx'
    )
#     import docx

# car = (3, 'Hyundai Creta', 'Hyundai', 'https://imgd.aeplcdn.com/664x374/n/cw/ec/41564/hyundai-creta-right-front-three-quarter9.jpeg?q=85', 'The Hyundai Creta has now entered its second generation and was first showcased at the 2020 Auto Expo. It’s the king of the D-segment and is offered in 14 variants across three engine and gearbox options.\xa0On the outside, it gets distinctive highlights in the form of three dimensional cascading grille with muscle like vertical and horizontal patterns and sculpted bumper with skid plates. The flared wheel arches with thoughtfully crafted crease lines enhance its overall appeal.\xa0The vehicle gets two boomerang shaped LED DRLs, LED Head Lamps and smartly positioned fog lamps to give the SUV a futuristic look. The new Creta gets class-leading features like panoramic sunroof, LED headlamps and enhanced Bluelink 1.5. The unveiled vehicle featured blacked-out windows and windscreens and more details about the interior will be known closer to its launch date.\xa0The Creta is offered with BS6 compliant petrol and diesel engine options. The standard petrol motor is a 1.5-litre producing 112bhp/144Nm and can be had with a CVT or a six-speed manual. The more exciting petrol motor is a 1.4-litre Turbo GDi unit producing 138bhp/242Nm and is only being offered with a seven-speed DCT. The sole diesel on offer is a 1.5-litre four-cylinder unit producing 113bhp/250Nm and can be had with a six-speed manual or a six-speed automatic. \xa0\xa0The Hyundai Creta is the most successful car in the segment since it was launched in 2015 and took the segment by a storm thanks to its combination of powerful engines and strong feature list. It’s a rival to the likes of the Kia Seltos, Nissan Kicks, Renault Duster and Captur as well as cars like the MG Hector and Tata Harrier.', 10.22, 17.93, 17.00, 21.00, True, True, True, True, False, False, 5)

# doc = docx.Document()
# head =doc.add_heading(car[1],0)
# head.alignment = 1
# doc.add_paragraph(str(car[4]))
# doc.add_paragraph('Price range: '+str(car[5])+' lakh to '+str(car[6])+' lakh')
# doc.add_paragraph('Mileage: '+str(car[7])+' kmpl '+'-'+str(car[8])+' kmpl')
# doc.add_paragraph('Seating capacity: '+str(car[15]))

# doc.save('sampletest.docx')


@app.route("/car/<int:id>")
def car(id):
    cur = dbcon.cursor()

    cur.execute("SELECT * FROM cars WHERE id=%s", (id,))

    car = cur.fetchone()
    cur.close()

    return render_template('Car.html', car=car)
   

@app.route("/results", methods=['GET', 'POST'])
def results():
    app.logger.info(request.form)
    if request.method == 'POST':
        company = request.form.get('company')

        price_starting = float(request.form.get('price_starting'))
        price_topend = float(request.form.get('price_topend'))

        manual = request.form.get('manual')
        if manual:
            manual = True

        automatic = request.form.get('automatic')
        if automatic:
            automatic = True

        petrol = request.form.get('petrol')
        if petrol:
            petrol = True

        diesel = request.form.get('diesel')
        if diesel:
            diesel = True

        cng = request.form.get('cng')
        if cng:
            cng = True

        electric = request.form.get('electric')
        if electric:
            cng = electric

        seating_capacity = int(request.form.get('seating_capacity'))

        car_sql = """
            SELECT id, name, company, image, price_starting, price_topend, mileage_l, mileage_u, manual, automatic, petrol, diesel, cng, electric, seating_capacity
            FROM cars
            WHERE company = %s
                AND price_starting > %s 
                AND price_topend < %s
                AND manual = COALESCE(%s, manual)
                AND automatic = COALESCE(%s, automatic)
                AND petrol = COALESCE(%s, petrol)
                AND diesel = COALESCE(%s, diesel)
                AND cng = COALESCE(%s, cng)
                AND electric = COALESCE(%s, electric)
                AND seating_capacity = %s;
                """

        car_params = (company, price_starting, price_topend, manual, automatic, petrol, diesel, cng, electric, seating_capacity)

        cur = dbcon.cursor()

        cur.execute(car_sql, car_params)

        cars = cur.fetchall()

        cur.execute("SELECT * FROM company WHERE name=%s", (company,))   

        company = cur.fetchone()

        app.logger.info(company)
        return render_template("Home.html", company=company,cars=cars)

if __name__ == "__main__":
    app.run(debug=True)
    dbcon.close()


